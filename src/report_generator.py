from docx import Document


class ReportGenerator:

    @staticmethod
    def create_report(
            auc,
            ks,
            gini,
            psi,
            report_path
    ):

        doc = Document()

        doc.add_heading(
            "Model Validation Report",
            level=1
        )

        doc.add_paragraph(
            f"AUC : {auc:.2f}"
        )

        doc.add_paragraph(
            f"KS : {ks:.2f}"
        )

        doc.add_paragraph(
            f"GINI : {gini:.2f}"
        )

        doc.add_paragraph(
            f"PSI : {psi:.2f}"
        )

        doc.save(report_path)